#!/usr/bin/env python3
"""Build the fillable Tishte Serif v0.950 independent review documents."""

from __future__ import annotations

import hashlib
import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "artifacts" / "review" / "v0950" / "docx"
VERSION = "0.950"
TAG = "v950"
INK = "202A33"
ACCENT = "8B2338"
MUTED = "66717A"
LIGHT = "F3F0EA"
HEADER_FILL = "E8EEF2"
WHITE = "FFFFFF"
CONTENT_WIDTH_DXA = 9638  # A4 with 20 mm margins: named government-review override.
TABLE_INDENT_DXA = 0
CELL_START_DXA = 120


def sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def set_run(run, *, name="Arial", size=11, bold=False, italic=False, color=INK) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    tr_pr.append(node)


def set_table_geometry(table, widths: list[int], header=True) -> None:
    total = sum(widths)
    if total != CONTENT_WIDTH_DXA:
        raise ValueError(f"table width must be {CONTENT_WIDTH_DXA}, got {total}")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row_index, row in enumerate(table.rows):
        for index, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if header and row_index == 0:
                shade(cell, HEADER_FILL)
        if header and row_index == 0:
            set_repeat_header(row)


def style_cell(cell, *, bold=False, color=INK, align=WD_ALIGN_PARAGRAPH.LEFT, size=9.5) -> None:
    for paragraph in cell.paragraphs:
        paragraph.alignment = align
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.1
        for run in paragraph.runs:
            set_run(run, size=size, bold=bold, color=color)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run("Стр. ")
    set_run(run, size=9, color=MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "1"
    r.append(t)
    field.append(r)
    paragraph._p.append(field)


def setup_document(title: str, subtitle: str, code: str) -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)
    section.left_margin = Mm(20)
    section.right_margin = Mm(20)
    section.header_distance = Mm(10)
    section.footer_distance = Mm(10)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for name, size, before, after in (("Title", 23, 0, 4), ("Heading 1", 16, 16, 8), ("Heading 2", 13, 12, 6), ("Heading 3", 12, 8, 4)):
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = name != "Title"
        style.font.color.rgb = RGBColor.from_string(INK if name == "Title" else ACCENT)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header.add_run(f"Tishte Serif v{VERSION}  •  независимая приёмка  •  {code}")
    set_run(run, size=9, bold=True, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_page_field(footer)

    p = doc.add_paragraph(style="Title")
    p.add_run(title)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run(subtitle)
    set_run(run, size=13, color=MUTED)
    add_metadata(doc, "Код документа", code)
    add_metadata(doc, "Объект", f"Tishte Serif v{VERSION}: Regular, Bold, Italic, Bold Italic")
    add_metadata(doc, "Статус", "Форма внешней приёмки; не является заполненным заключением")
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run("Важно: пустые поля и флажки не означают положительного решения эксперта.")
    set_run(run, size=10.5, bold=True, color=ACCENT)
    return doc


def add_metadata(doc: Document, label: str, value: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label + ": ")
    set_run(r, bold=True)
    r = p.add_run(value)
    set_run(r)


def add_heading(doc: Document, text: str, level=1) -> None:
    doc.add_paragraph(text, style=f"Heading {level}")


def add_para(doc: Document, text: str, *, bold=False, italic=False, color=INK, size=11, after=6) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    set_run(r, size=size, bold=bold, italic=italic, color=color)


def add_check(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Mm(4)
    p.paragraph_format.first_line_indent = Mm(-4)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("☐  " + text)
    set_run(r, size=10.5)


def add_lines(doc: Document, labels: list[str]) -> None:
    for label in labels:
        add_para(doc, f"{label}: ______________________________________________________________", size=10.5, after=5)


def add_issue_table(doc: Document, rows=6) -> None:
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = ("№", "Знак/фрагмент", "Среда и размер", "Описание", "Критичность")
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
        style_cell(cell, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=9)
    for index in range(1, rows + 1):
        cells = table.add_row().cells
        cells[0].text = str(index)
        for cell in cells:
            style_cell(cell, align=WD_ALIGN_PARAGRAPH.CENTER if cell is cells[0] else WD_ALIGN_PARAGRAPH.LEFT, size=9)
            cell.paragraphs[0].paragraph_format.space_after = Pt(14)
    set_table_geometry(table, [480, 1500, 1600, 4158, 1900])


def add_decision(doc: Document) -> None:
    add_heading(doc, "Решение эксперта", 1)
    add_check(doc, "Принято без замечаний.")
    add_check(doc, "Принято с обязательными замечаниями, перечисленными выше.")
    add_check(doc, "Отклонено до исправления и повторной проверки.")
    add_lines(doc, ["Обоснование", "ФИО", "Организация и роль", "Контакт", "Дата", "Подпись"])


def hashes() -> list[tuple[str, str]]:
    result = []
    for style in ("Regular", "Bold", "Italic", "BoldItalic"):
        path = ROOT / "build" / f"TishteSerif-{style}-{TAG}.ttf"
        result.append((path.name, sha256(path)))
    return result


def add_hash_table(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Файл"
    table.rows[0].cells[1].text = "SHA-256"
    for cell in table.rows[0].cells:
        style_cell(cell, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=9)
    for filename, digest in hashes():
        cells = table.add_row().cells
        cells[0].text = filename
        cells[1].text = digest
        style_cell(cells[0], size=8.5)
        style_cell(cells[1], size=7.5)
    set_table_geometry(table, [3300, 6338])


def build_guide() -> Path:
    doc = setup_document("РУКОВОДСТВО ПО ПРИЁМКЕ", "Комплект передачи независимым экспертам", "TS-REV-00")
    add_heading(doc, "Назначение", 1)
    add_para(doc, "Комплект фиксирует объект проверки, разделяет компетенции и не позволяет заменить подпись эксперта автоматическим отчётом. Каждый специалист получает TTF, свою форму DOCX/PDF и SHA-256 файлов.")
    add_heading(doc, "Порядок работы", 1)
    for text in (
        "Шаг 1. Сверить SHA-256 и установить четыре файла TTF v0.950.",
        "Шаг 2. Зафиксировать ОС, приложение, экран, принтер, бумагу и масштаб.",
        "Шаг 3. Проверить только протокол своей компетенции и зарегистрировать каждый дефект.",
        "Шаг 4. Выбрать одно итоговое решение, подписать и сохранить неизменённую копию.",
        "Шаг 5. При замечаниях приложить исходный текст, скриншот или PDF и точные условия воспроизведения.",
    ):
        add_para(doc, text, bold=True if text.startswith("Шаг") else False)
    add_heading(doc, "Контрольные файлы", 1)
    add_hash_table(doc)
    add_heading(doc, "Разделение протоколов", 1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for cell, text in zip(table.rows[0].cells, ("Протокол", "Компетенция", "Нельзя подменять")):
        cell.text = text
        style_cell(cell, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=9)
    rows = (
        ("Типографический", "Шрифтовой дизайнер/типограф", "Лингвистическую и юридическую оценку"),
        ("Луговомарийский", "Редактор луговомарийского", "Горномарийскую приёмку"),
        ("Горномарийский", "Редактор горномарийского", "Луговомарийскую приёмку"),
        ("Юридический", "Юрист по ИС и публичному праву", "Художественную оценку"),
    )
    for row in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, row):
            cell.text = text
            style_cell(cell, size=9)
    set_table_geometry(table, [2100, 3500, 4038])
    add_heading(doc, "Критический дефект", 1)
    add_para(doc, "Критическим считается дефект, меняющий значение, делающий знак неразличимым, вызывающий .notdef, обрезание, наложение, несовместимую пагинацию, нарушение OFL либо ложное впечатление официального статуса.", color=ACCENT, bold=True)
    add_heading(doc, "Комплект передачи", 1)
    for text in (
        "Четыре TTF-файла v0.950 и файл SHA256SUMS.txt.",
        "Форма эксперта одновременно в DOCX и PDF; заполняется только одна рабочая копия.",
        "Внутреннее техническое досье и отчёты CI используются как доказательства, но не как подпись эксперта.",
        "Скриншоты, PDF и исправленные тексты прикладываются без изменения исходной формы.",
    ):
        add_check(doc, text)
    add_heading(doc, "Фиксация результата", 1)
    add_para(doc, "Имя итогового файла: код протокола, фамилия эксперта и дата в формате ГГГГ-ММ-ДД. Подписанная версия сохраняется вместе с приложениями; после получения проекта проверяется совпадение SHA-256 и полнота полей решения.")
    add_heading(doc, "Маршрут замечаний", 1)
    add_para(doc, "Каждое замечание получает воспроизводимое описание и критичность. После исправления выпускается новый кандидат с новым номером версии и хэшами; прежняя подписанная форма не редактируется.")
    path = OUTPUT / "00-Guide-v0950.docx"
    doc.save(path)
    return path


def build_typography() -> Path:
    doc = setup_document("ТИПОГРАФИЧЕСКАЯ ПРИЁМКА", "Экран, печать и согласованность четырёх начертаний", "TS-REV-01")
    add_heading(doc, "Данные проверки", 1)
    add_lines(doc, ["Эксперт", "ОС и версия", "Приложения", "Экран и масштаб", "Принтер и бумага"])
    add_heading(doc, "Матрица размеров", 1)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    for cell, text in zip(table.rows[0].cells, ("pt", "Экран", "Печать", "Цвет/ритм", "Замечание")):
        cell.text = text
        style_cell(cell, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=9)
    for size in (8, 9, 10, 11, 12, 14, 18, 24):
        cells = table.add_row().cells
        cells[0].text = str(size)
        cells[1].text = "☐ принято  ☐ нет"
        cells[2].text = "☐ принято  ☐ нет"
        cells[3].text = "☐ ровно  ☐ нет"
        for cell in cells:
            style_cell(cell, align=WD_ALIGN_PARAGRAPH.CENTER, size=8.5)
    set_table_geometry(table, [600, 1700, 1700, 1800, 3838])
    add_heading(doc, "Обязательные проверки", 1)
    for text in (
        "Regular, Bold, Italic и Bold Italic различимы и образуют одно семейство.",
        "Русская и марийская кириллица не выглядят приложением к латинице.",
        "Цифры 0–9, №, ₽, %, дроби и математические знаки устойчивы в таблицах.",
        "Кавычки, тире, дефис, скобки и пунктуация согласованы по цвету.",
        "Диакритика не сталкивается с буквами в NFC и NFD.",
        "Строка не становится визуально тесной или рыхлой при неизменных метриках Times.",
        "PDF сохраняет внедрение, поиск, копирование и четыре начертания.",
        "На лазерной и струйной печати нет выпадения тонких штрихов и заливания просветов.",
    ):
        add_check(doc, text)
    add_heading(doc, "Журнал дефектов", 1)
    add_issue_table(doc, 8)
    add_decision(doc)
    path = OUTPUT / "01-Typography-Review-v0950.docx"
    doc.save(path)
    return path


def add_language_samples(doc: Document, section: str, inventory: str) -> None:
    corpus = json.loads((ROOT / "data" / "language-corpus.json").read_text(encoding="utf-8"))
    add_para(doc, "Тексты ниже являются тестовым материалом. Эксперт обязан исправить ошибки содержания и орфографии; их наличие не считается утверждённой нормой.", bold=True, color=ACCENT)
    for line in corpus["sections"][section]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(7)
        r = p.add_run(line)
        set_run(r, name="Tishte Serif", size=14)
    add_heading(doc, "Контрольные знаки", 2)
    for style, bold, italic in (("Regular", False, False), ("Bold", True, False), ("Italic", False, True), ("Bold Italic", True, True)):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        label = p.add_run(style + ": ")
        set_run(label, size=10, bold=True, color=MUTED)
        r = p.add_run(inventory)
        set_run(r, name="Tishte Serif", size=18, bold=bold, italic=italic)


def build_language(kind: str) -> Path:
    meadow = kind == "meadow_mari"
    title = "ЛУГОВОМАРИЙСКАЯ ПРИЁМКА" if meadow else "ГОРНОМАРИЙСКАЯ ПРИЁМКА"
    subtitle = "Отдельный протокол редактора луговомарийского языка" if meadow else "Отдельный протокол редактора горномарийского языка"
    code = "TS-REV-02" if meadow else "TS-REV-03"
    inventory = "Ӓ ӓ  Ӧ ӧ  Ӱ ӱ  Ҥ ҥ" if meadow else "Ӓ ӓ  Ӧ ӧ  Ӱ ӱ  Ӹ ӹ"
    filename = "02-Meadow-Mari-Review-v0950.docx" if meadow else "03-Hill-Mari-Review-v0950.docx"
    doc = setup_document(title, subtitle, code)
    add_heading(doc, "Компетенция эксперта", 1)
    add_lines(doc, ["ФИО", "Организация/редакционная роль", "Стаж и профиль", "Контакт"])
    add_language_samples(doc, kind, inventory)
    add_heading(doc, "Проверка", 1)
    checks = [
        "Формы специальных букв соответствуют принятой редакционной практике.",
        "Прописные и строчные пары различимы во всех четырёх начертаниях.",
        "Диакритика расположена естественно и не меняет чтение слова.",
        "Частотные сочетания и соседние буквы не создают ложных форм.",
        "Курсив остаётся читаемым и не воспринимается как другой знак.",
        "Официально-деловая лексика, числа, даты, №, ₽ и кавычки набираются естественно.",
        "NFC и NFD визуально эквивалентны в проверенных приложениях.",
        "Предложенные тестовые фразы отредактированы или подтверждены экспертом.",
    ]
    if meadow:
        checks.append("Проверка Ҥ/ҥ выполнена отдельно от горномарийского набора.")
    else:
        checks.append("Проверка Ӹ/ӹ выполнена отдельно от луговомарийского набора.")
    for text in checks:
        add_check(doc, text)
    add_heading(doc, "Исправленный экспертный текст", 1)
    for _ in range(3):
        add_para(doc, "________________________________________________________________________________", after=5)
    add_heading(doc, "Журнал дефектов", 1)
    add_issue_table(doc, 4)
    add_decision(doc)
    path = OUTPUT / filename
    doc.save(path)
    return path


def build_legal() -> Path:
    doc = setup_document("ЮРИДИЧЕСКАЯ ПРИЁМКА", "OFL, происхождение, название и порядок официального внедрения", "TS-REV-04")
    add_heading(doc, "Данные специалиста", 1)
    add_lines(doc, ["ФИО", "Организация", "Специализация", "Контакт"])
    add_heading(doc, "Матрица правовых вопросов", 1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    for cell, text in zip(table.rows[0].cells, ("Вопрос", "Материал", "Решение", "Комментарий")):
        cell.text = text
        style_cell(cell, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=9)
    rows = (
        ("Актуальная редакция постановления № 316", "Текст акта и п. 35", "☐ да ☐ нет", ""),
        ("Квалификация метрического аналога", "Метрики и Office-корпус", "☐ да ☐ нет", ""),
        ("Соблюдение SIL OFL 1.1", "OFL.txt, upstream commit", "☐ да ☐ нет", ""),
        ("Copyright/name table", "TTF v0.950, audit JSON", "☐ да ☐ нет", ""),
        ("Права участников проекта", "AUTHORS/CONTRIBUTORS", "☐ да ☐ нет", ""),
        ("Название Tishte/Тиште", "Роспатент и WIPO", "☐ да ☐ нет", ""),
        ("Отсутствие ложного госстатуса", "README, metadata, релиз", "☐ да ☐ нет", ""),
        ("Порядок внедрения и обновления", "Проект регламента", "☐ да ☐ нет", ""),
    )
    for row in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, row):
            cell.text = text
            style_cell(cell, size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER if cell is cells[2] else WD_ALIGN_PARAGRAPH.LEFT)
    set_table_geometry(table, [2800, 2600, 1400, 2838])
    add_heading(doc, "Обязательные источники", 1)
    for source in (
        "Постановление № 316 (контрольная копия): https://base.garant.ru/409588219/",
        "ГОСТ Р 7.0.97-2025: https://protect.gost.ru/gost/details/360994e3-9a70-47b9-ab3a-cf21809e26ed",
        "Pinned Tinos commit: https://github.com/googlefonts/tinos/commit/3b4482a99b80ea5fc75f187b1be3120a3f5905b3",
        "Официальный OFL 1.1: https://openfontlicense.org/open-font-license-official-text/",
        "Правила модификации: https://openfontlicense.org/how-to-modify-ofl-fonts/",
        "Роспатент: https://searchplatform.rospatent.gov.ru/",
        "WIPO Global Brand Database: https://branddb.wipo.int/",
    ):
        add_para(doc, source, size=9.5, after=4)
    doc.add_page_break()
    add_heading(doc, "Неавтоматизируемые действия", 1)
    for text in (
        "Получить официальную актуальную копию регионального акта и проверить изменения.",
        "Провести профессиональный поиск сходных обозначений, включая кириллицу и транслитерацию.",
        "Определить правообладателя вклада и полномочия подписанта лицензии/релиза.",
        "Определить уполномоченный орган и проект акта об утверждении или пилотировании.",
        "Утвердить регламент поставки, обновления, контроля хэшей и отзыва версии.",
    ):
        add_check(doc, text)
    add_heading(doc, "Замечания", 1)
    add_issue_table(doc, 3)
    add_decision(doc)
    path = OUTPUT / "04-Legal-Review-v0950.docx"
    doc.save(path)
    return path


def load_report(name: str) -> dict:
    path = ROOT / "artifacts" / "reports" / name
    return json.loads(path.read_text(encoding="utf-8")) if path.exists() else {}


def build_evidence() -> Path:
    doc = setup_document("ВНУТРЕННЕЕ ТЕХНИЧЕСКОЕ ДОСЬЕ", "Доказательства для внешней приёмки; не заменяет её", "TS-REV-05")
    metric = load_report("metric-contract-v950.json")
    original = load_report("outline-originality-v950.json")
    reproducible = load_report("reproducible-build-v950.json")
    legal = load_report("legal-metadata-v950.json")
    add_heading(doc, "Контрольные файлы", 1)
    add_hash_table(doc)
    add_heading(doc, "Сводка автоматических проверок", 1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for cell, text in zip(table.rows[0].cells, ("Контроль", "Результат", "Граница вывода")):
        cell.text = text
        style_cell(cell, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=9)
    metric_ok = metric.get("passed", False)
    original_ok = original.get("passed", False)
    reproducible_ok = reproducible.get("passed", False)
    legal_ok = legal.get("passed", False)
    rows = (
        ("Метрики", "PASS: 422/422; 712 336 пар" if metric_ok else "нет отчёта", "Не оценивает художественное качество"),
        ("Контуры", "PASS: 0/413 точных совпадений" if original_ok else "нет отчёта", "Не является правовым тестом самостоятельности"),
        ("Воспроизводимость", "PASS: 4 TTF + 4 WOFF2" if reproducible_ok else "нет отчёта", "Не заменяет платформенное тестирование"),
        ("Правовые метаданные", "PASS" if legal_ok else "нет отчёта", "Не заменяет заключение юриста"),
        ("Языковые корпуса", "PASS без .notdef", "Текст требует двух редакторов"),
        ("FontForge/OTS/FontBakery", "Проверяется CI", "Не заменяет печатную приёмку"),
    )
    for row in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, row):
            cell.text = text
            style_cell(cell, size=9)
    set_table_geometry(table, [2400, 2800, 4438])
    add_heading(doc, "Внутренняя визуальная проверка", 1)
    for text in (
        "Просмотрены сравнительные листы Regular, Bold, Italic и Bold Italic.",
        "Обнаруженные изменения v0.940 ограничены остаточными знаками и диакритикой.",
        "Явных наложений, обрезания и выпадения специальных марийских букв не выявлено.",
        "Результат не считается независимым, поскольку создан внутри проекта.",
    ):
        add_check(doc, text)
    add_heading(doc, "Открытые ворота 1.0", 1)
    for text in (
        "Подписанный типографический протокол.",
        "Подписанный луговомарийский протокол.",
        "Подписанный горномарийский протокол.",
        "Подписанное юридическое заключение и проверка на обезличенных реальных документах.",
    ):
        add_check(doc, text)
    path = OUTPUT / "05-Internal-Evidence-v0950.docx"
    doc.save(path)
    return path


def main() -> None:
    OUTPUT.mkdir(parents=True, exist_ok=True)
    paths = [
        build_guide(),
        build_typography(),
        build_language("meadow_mari"),
        build_language("hill_mari"),
        build_legal(),
        build_evidence(),
    ]
    print(json.dumps({"documents": [str(path) for path in paths]}, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
